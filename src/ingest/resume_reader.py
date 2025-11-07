"""Utilities for ingesting resume files into normalized text."""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

logger = logging.getLogger(__name__)


@dataclass
class ResumeDocument:
    """Container for resume text and derived sections."""

    text: str
    sections: dict[str, str]
    bullets: list[str]


_SECTION_HEADERS = (
    "experience",
    "work experience",
    "education",
    "skills",
    "projects",
    "summary",
)


def _read_pdf(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text  # type: ignore

        return extract_text(path) or ""
    except Exception as exc:  # pragma: no cover - optional dependency path
        logger.warning("Falling back to text extraction for %s: %s", path, exc)
        try:
            return path.read_text(encoding="utf-8")
        except Exception:
            return ""


def _read_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError("python-docx is required to read DOCX resumes") from exc

    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _clean_lines(lines: Iterable[str]) -> list[str]:
    cleaned: list[str] = []
    for raw in lines:
        line = raw.replace("\u2022", "-")
        line = re.sub(r"\s+", " ", line).strip()
        if line:
            cleaned.append(line)
    return cleaned


def _split_sections(text: str) -> dict[str, str]:
    sections: dict[str, str] = {}
    pattern = re.compile(r"^(?P<header>[A-Za-z ]{3,}):?$", re.IGNORECASE)
    current_header = "general"
    buffer: list[str] = []
    for line in text.splitlines():
        match = pattern.match(line.strip())
        if match and match.group("header").lower() in _SECTION_HEADERS:
            if buffer:
                sections[current_header] = "\n".join(buffer).strip()
                buffer = []
            current_header = match.group("header").lower()
        else:
            buffer.append(line)
    if buffer:
        sections[current_header] = "\n".join(buffer).strip()
    return sections


def _extract_bullets(text: str) -> list[str]:
    bullet_pattern = re.compile(r"^(?:[-*\u2022]|\d+\.)\s+(.*)")
    bullets: list[str] = []
    for line in text.splitlines():
        match = bullet_pattern.match(line.strip())
        if match:
            bullets.append(match.group(1).strip())
    if not bullets:
        bullets = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
    return bullets


def read_resume(path: str | Path) -> ResumeDocument:
    """Read a resume from the given path and produce normalized text."""

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix in {".doc", ".docx"}:
        text = _read_docx(path)
    else:
        text = _read_text(path)

    lines = _clean_lines(text.splitlines())
    normalized = "\n".join(lines)
    sections = _split_sections(normalized)
    bullets = _extract_bullets(normalized)
    logger.debug("Loaded resume %s with %d bullets", path, len(bullets))
    return ResumeDocument(text=normalized, sections=sections, bullets=bullets)
